import docx
def findParagraphIndex(doc, currentIndex, paraType, paraText = 'DontCare'):
    for i in range(currentIndex, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name == paraType and paraText == 'DontCare':
            return i
            
        if doc.paragraphs[i].style.name == paraType and doc.paragraphs[i].text == paraText:
            return i
            
        return -1 
            

def buildFromTemplate(templateFile):
    doc = docx.Document(templateFile)
    textReplace = 'newText'
    paraIndex = 0
   
    while True:
   
        paraIndex = findParagraphIndex(doc, paraIndex, 'Heading 1')
        if (paraIndex == -1):
            break
        doc.paragraphs[paraIndex].text = textReplace
    
    return doc
    
    
doc = docx.Document()
doc = buildFromTemplate('Resume.docx')
doc.save('new.docx')
